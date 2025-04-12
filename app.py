from dotenv import load_dotenv
load_dotenv()
import logging
from flask import Flask, request, jsonify
import requests
from ai21 import AI21Client
from ai21.models.chat import UserMessage
from docx import Document
import os
import re
import time
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import urllib.parse

app = Flask(__name__)

# Configuração de logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Manipulador de console
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)

# Formato dos logs
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
console_handler.setFormatter(formatter)

# Adiciona o manipulador ao logger
logger.addHandler(console_handler)

# Teste
logger.info('Teste de logging no console')

# Chaves de API
AI21_API_KEY = os.getenv("AI21_API_KEY")
AI21_API_KEY_TRABALHO = os.getenv("AI21_API_KEY_TRABALHO")
WHATSAPP_ACCESS_TOKEN = os.getenv("WHATSAPP_ACCESS_TOKEN")

# Configurações
HUGGINGFACE_API_URL = "https://api-inference.huggingface.co/models/black-forest-labs/FLUX.1-dev"
HUGGINGFACE_API_KEY = os.getenv("HUGGINGFACE_API_KEY")
WHATSAPP_API_URL = "https://graph.facebook.com/v22.0"
TEMP_DIR = r'C:/tmp'

# Inicializando os clientes AI21
client_perguntas = AI21Client(api_key=AI21_API_KEY)
client_trabalhos = AI21Client(api_key=AI21_API_KEY_TRABALHO)

# Variável para armazenar o estado da conversa de cada usuário
user_state = {}

# Registro para garantir que o arquivo DOCX não seja enviado mais de uma vez
file_sent = {}






# Função para gerar imagens com Hugging Face
def generate_images_with_huggingface(text_description, num_images=5):
    headers = {
        "Authorization": f"Bearer {HUGGINGFACE_API_KEY}",
        "Content-Type": "application/json"
    }
    
    image_paths = []
    
    for i in range(num_images):
        data = {"inputs": f"{text_description}"}
        
        try:
            response = requests.post(HUGGINGFACE_API_URL, headers=headers, json=data)
            response.raise_for_status()
            image_data = response.content

            # Sanitiza a descrição para o nome do arquivo
            sanitized_description = urllib.parse.quote_plus(text_description[:50])  # Limita o tamanho e sanitiza
            image_filename = f"{sanitized_description}_Nexusbot_{i+1}.png"
            image_path = os.path.join(TEMP_DIR, image_filename)

            with open(image_path, 'wb') as f:
                f.write(image_data)

            image_paths.append(image_path)
        except requests.RequestException as e:
            logger.error(f"Erro ao gerar imagem {i+1}: {e}")

    return image_paths[:3]  # Garantir que no máximo 3 imagens sejam retornadas


# Função para fazer upload da imagem para o WhatsApp
def upload_image_to_whatsapp(image_path, phone_number_id, token):
    upload_url = f"{WHATSAPP_API_URL}/{phone_number_id}/media"
    headers = {"Authorization": f"Bearer {token}"}

    if not os.path.exists(image_path):
        logger.error(f"Arquivo não encontrado: {image_path}")
        return None
    
    try:
        with open(image_path, 'rb') as file:
            files = {
                'file': ('output_image.png', file, 'image/png'),
                'messaging_product': (None, 'whatsapp')
            }
            media_response = requests.post(upload_url, headers=headers, files=files)
            media_response.raise_for_status()
            return media_response.json().get('id')
    except requests.RequestException as e:
        logger.error(f"Erro ao enviar o arquivo para o WhatsApp: {e}")
        return None

# Função para enviar imagem via WhatsApp
def send_image_via_whatsapp(phone_number, media_id, token, phone_number_id):
    message_url = f"{WHATSAPP_API_URL}/{phone_number_id}/messages"
    data = {
        "messaging_product": "whatsapp",
        "to": phone_number,
        "type": "image",
        "image": {
            "id": media_id,
            "caption": "Aqui está a imagem gerada🎉!"
        }
    }

    try:
        response = requests.post(message_url, headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"}, json=data)
        response.raise_for_status()
        logger.info(f"Imagem enviada com sucesso para {phone_number}!")
    except requests.RequestException as e:
        logger.error(f"Erro ao enviar imagem: {e}")


# Lista de palavras-chave restritas
RESTRICTED_KEYWORDS = [
    "vagina", "pênis", "sexo", "pornografia", "adulto", "nudez", "obsceno", "explícito","porn","xxvideo","xnxx","+18","seios","bordel","sem roupa",
    "drogas", "violência", "discriminação", "abuso", "menor", "criança", "Novinha nua","sex", "porn", "nude", "erotic", "rape", "abuse", "violence", "harassment", "adult", "fetish", "nua", "conteudo +18", "conteudo +18","conteúdo +18","conteúdo +18","conteúdo adulto"
]

# Função para verificar se a descrição contém palavras restritas
def contains_restricted_keywords(description):
    description_lower = description.lower()
    for keyword in RESTRICTED_KEYWORDS:
        if keyword in description_lower:
            return True
    return False


# Função principal para lidar com a geração e envio de imagens
def handle_image_generation_and_sending(phone_number, user_message, phone_number_id, token):

     # Verifica se a mensagem do usuário contém palavras restritas
    if contains_restricted_keywords(user_message):
        send_whatsapp_message(phone_number, "🚫 Desculpe, não podemos gerar imagens com conteúdos impróprios ou sensíveis.")
        return

    # Enviar mensagem inicial de paciência
    send_whatsapp_message(phone_number, "🎉 Estamos gerando a imagem. Por favor, tenha paciência, isso pode levar alguns momentos ")
    

    # Gerar imagens
    image_paths = generate_images_with_huggingface(user_message)
    
    all_images_sent = True  # Variável para monitorar o sucesso do envio

    # Enviar imagens
    for image_path in image_paths:
        media_id = upload_image_to_whatsapp(image_path, phone_number_id, token)
        if media_id:
            send_image_via_whatsapp(phone_number, media_id, token, phone_number_id)
        else:
            all_images_sent = False  # Caso alguma imagem não seja enviada corretamente

    # Após o envio das imagens, verificar se todas foram enviadas
    if all_images_sent:
        logger.info(f"Todas as imagens foram enviadas para {phone_number}.")
    else:
        logger.error(f"Houve um problema ao enviar algumas imagens para {phone_number}.")



# Função para realizar chamadas à API AI21 com retry e timeout
def chamar_api_ai21(client, model, messages, max_retries=3, delay=5, timeout=60):
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(model=model, messages=messages, top_p=1.0, timeout=timeout)
            return response
        except Exception as e:
            logger.error(f"Erro ao conectar com a API AI21: {e}")
            if attempt < max_retries - 1:
                time.sleep(delay)
            else:
                raise

def perguntar_ai21(mensagem):
    try:
        prompt = (
            "Se o usuário fizer perguntas  sobre o criador do chatbot ou sobre sua origem e proveniência, responda: "
            "'Eu sou o NexusBot, um chatbot integrado ao WhatsApp, criado pelo estudante *Jacinto Sérgio Robate*, que estuda na `Universidade Católica de Moçambique` e está se formando no curso de `Tecnologias de Informação`*(IT)*.' "
            "Para qualquer outra pergunta, responda de forma objetiva e explicativa, sempre proporcionando informações úteis e claras ao usuário."
            "Se a pergunta não for clara informe ao usuário que ele deve ser mais claro em sua pergunta. "
            f"{mensagem}\n\nResposta:"
        )

   
        messages = [UserMessage(content=prompt)]
        response = chamar_api_ai21(client_perguntas, "jamba-1.5-large", messages)
        
        resposta = response.choices[0].message.content.strip()
        if not resposta:
            resposta = "Desculpe, não consegui gerar uma resposta adequada. Tente novamente mais tarde."
        resposta_formatada = format_response(resposta)
        return resposta_formatada
    except Exception as e:
        logger.error(f"Erro ao conectar com a API AI21: {e}")
        return format_response(f"ERRO AO CONECTAR COM A API AI21: {str(e)}")

def gerar_trabalho_ai21(tema):
    try:
        prompt = f"Você é um assistente especializado em gerar trabalhos acadêmicos completos. Por favor, escreva um trabalho academico bem estruturado e detalhado com todos os subtopicos e subtemas que julgares importantes e necessarios sobre o seguinte tema:\n\nTema: {tema}\n\nTrabalho:"
        messages = [UserMessage(content=prompt)]
        response = chamar_api_ai21(client_trabalhos, "jamba-1.5-large", messages)
        
        trabalho = response.choices[0].message.content.strip()
        if not trabalho:
            trabalho = "Desculpe, não consegui gerar um trabalho adequado sobre este tema. Tente novamente."
        return trabalho
    except Exception as e:
        logger.error(f"Erro ao conectar com a API AI21: {e}")
        return f"ERRO AO CONECTAR COM A API AI21: {str(e)}"





def criar_trabalho_academico(tema, conteudo):
    try:
        doc = Document()

        # Adiciona o "Tema" em negrito, centralizado, com tamanho 13.5
        tema_paragrafo = doc.add_paragraph()
        tema_paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o tema
        tema_run = tema_paragrafo.add_run(f'Tema: {tema}')
        tema_run.bold = True
        tema_run.font.name = 'Times New Roman'
        tema_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        tema_run.font.size = Pt(13.5)

        # Remover caracteres indesejados do conteúdo
        conteudo_limpo = re.sub(r'[#+\*\-\'\"=]+', '', conteudo)

        # Dividimos o conteúdo em linhas
        linhas = conteudo_limpo.split('\n')

        for linha in linhas:
            linha_strip = linha.strip()

            if not linha_strip:
                continue  # Ignora linhas vazias

            # Verifica se a linha é um subtópico com numeração hierárquica (ex: 1.1, 2.2)
            if re.match(r'^\d+\.\d+\s', linha_strip):
                # Inicia a lista numerada para subtópicos com numeração hierárquica
                paragrafo = doc.add_paragraph(linha_strip, style='List Number')
                paragrafo.paragraph_format.left_indent = Pt(36)  # Recuo para subtópicos numerados
                run = paragrafo.runs[0]
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13.5)
            # Verifica se a linha é um subtópico normal (começa com número ou letra seguida por ponto)
            elif re.match(r'^\d+\.\s|^[a-zA-Z]+\.\s', linha_strip):
                paragrafo = doc.add_paragraph()
                paragrafo.paragraph_format.left_indent = Pt(36)  # Recuo do subtópico
                run = paragrafo.add_run(linha_strip)  # Remove os espaços em excesso
                run.bold = True  # Aplica o negrito ao subtópico
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # Garantir compatibilidade
                run.font.size = Pt(13.5)
            else:
                # Se não for subtópico, adiciona o conteúdo justificado
                paragrafo = doc.add_paragraph(linha_strip)
                paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justifica o conteúdo
                run = paragrafo.runs[0] if paragrafo.runs else paragrafo.add_run()  # Verifica se o parágrafo tem runs
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(12)

        # Sanitiza o nome do tema para gerar o nome do arquivo
        tema_sanitizado = re.sub(r'[\/\\\:\*\?\"<>\|]', '_', tema)
        tema_sanitizado = re.sub(r'\s+', ' ', tema_sanitizado).strip()
        file_name = f'NexusBOT_{tema_sanitizado}_.docx'
        file_path = os.path.join('/tmp', file_name)
        doc.save(file_path)
        return file_path, file_name
    except Exception as e:
        print(f"Erro ao criar o trabalho acadêmico: {e}")
        return None, None



def enviar_arquivo_whatsapp(phone_number, file_path, file_name):
    logger.info(f"Enviando arquivo para o número: {phone_number}, arquivo: {file_path}")
    phone_number_id = '678628148657980'
    upload_url = f"https://graph.facebook.com/v22.0/{phone_number_id}/media"
    headers = {
        "Authorization": f"Bearer {WHATSAPP_ACCESS_TOKEN}",
    }

    if not os.path.exists(file_path):
        logger.error(f"Arquivo não encontrado: {file_path}")
        return None
    
    with open(file_path, 'rb') as file:
        files = {
            'file': (file_name, file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            'messaging_product': (None, 'whatsapp')
        }
        media_response = requests.post(upload_url, headers=headers, files=files)

    if media_response.status_code != 200:
        logger.error(f"Erro ao enviar o arquivo para o WhatsApp: {media_response.status_code}, {media_response.text}")
        return None

    media_id = media_response.json().get('id')

    message_url = f"https://graph.facebook.com/v22.0/{phone_number_id}/messages"
    data = {
        "messaging_product": "whatsapp",
        "to": phone_number,
        "type": "document",
        "document": {
            "id": media_id,
            "filename": file_name
        }
    }
    response = requests.post(message_url, headers=headers, json=data)
    logger.info(f"Envio de arquivo: {response.status_code}, {response.text}")
    
    return response.json()

def enviar_mensagem_atualizacao(phone_number, index):
    mensagens_atualizacao = [
        "⏳ `Estamos trabalhando no seu pedido!` 📚\n\nPor favor, tenha paciência. Seu trabalho está quase pronto!",
        "🔄 `Estamos finalizando os detalhes do seu trabalho!` 💻\n\nAgradecemos pela paciência, está quase pronto!",
        "🕰️ `Estamos quase lá!` 🌟\n\nSó mais um pouco e seu trabalho estará pronto para você!",
        "🚀 `Preparando os últimos ajustes!` 📄\n\nSeu trabalho está a caminho, obrigado por aguardar!",
        "⏲️ `Estamos na reta final!` 🎯\n\nEm breve, seu trabalho estará disponível para você!",
    ]
    mensagem = mensagens_atualizacao[index]
    send_whatsapp_message(phone_number, mensagem)

# Expressão regular para capturar saudações como "oi", "ola", "olá", e variações com emojis
saudacoes_regex = re.compile(r"^(ola|olá|oi|cmkie|kmk|jacinto|bom dia|boa tarde|boa noite|puto|👋|)$", re.IGNORECASE)




def send_whatsapp_message(phone_number, message, buttons=None):
    phone_number_id = '678628148657980'
    url = f"https://graph.facebook.com/v22.0/{phone_number_id}/messages"
    headers = {
        "Authorization": f"Bearer {WHATSAPP_ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }

    data = {
        "messaging_product": "whatsapp",
        "to": phone_number,
        "type": "text",
        "text": {
            "body": message
        }
    }

    if buttons:
        data["type"] = "interactive"
        data["interactive"] = {
            "type": "button",
            "body": {
                "text": message
            },
            "action": {
                "buttons": buttons
            }
        }

    response = requests.post(url, headers=headers, json=data)
    logger.info(f"Envio de mensagem: {response.status_code}, {response.text}")
    return response.json()

def send_creator_info(phone_number):
    phone_number_id = '678628148657980'
    url = f"https://graph.facebook.com/v22.0/{phone_number_id}/messages"
    
    headers = {
        "Authorization": f"Bearer {WHATSAPP_ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }
    
    creator_message = (
        "🌟 Olá! Eu sou o NexusBot, criado por `Jacinto SR`, um estudante de TI com uma paixão por tecnologia e IA, Ele me criou para ajudar a todos. 🚀\n\n"
        "💡 O nome `NexusBot` vem de nexus, que significa conexão. Fui criado para conectar você com respostas rápidas e precisas, facilitando seu dia a dia!\n\n"
        "🛠️ Sou capaz de responder perguntas, gerar trabalhos acadêmicos, e muito mais. Estou aqui para ajudar você!\n\n"
        "🔗 Conecte-se na rede social abaixo:\n\n"
    )
    
    data = {
        "messaging_product": "whatsapp",
        "recipient_type": "individual",
        "to": phone_number,
        "type": "interactive",
        "interactive": {
            "type": "cta_url",
            "body": {
                "text": creator_message
            },
            "action": {
                "name": "cta_url",
                "parameters": {
                    "display_text": "Acessar LinkedIn",
                    "url": "https://www.linkedin.com/in/jacinto-robate-942a62267/"
                }
            }
        }
    }
    
    
    response = requests.post(url, headers=headers, json=data)
    logger.info(f"Envio de informações do criador: {response.status_code}, {response.text}")

    time.sleep(1)
    send_menu(phone_number)
    return response.json()

def format_response(text):
    return text



# Envia o menu interativo
def send_menu(phone_number):
    phone_number_id = '678628148657980'
    url = f"https://graph.facebook.com/v22.0/{phone_number_id}/messages"
    headers = {
        "Authorization": f"Bearer {WHATSAPP_ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }

    data = {
        "messaging_product": "whatsapp",
        "to": phone_number,
        "type": "interactive",
        "interactive": {
            "type": "list",
            "header": {
                "type": "text",
                "text": "👋NexusBot"
            },
            "body": {
                "text": "Como posso ajudar você hoje?✋😝🤚:"
            },
            "footer": {
                "text": "Selecione uma das opções abaixo"
            },
            "action": {
                "button": "Abrir opções ",
                "sections": [
                    {
                        "title": "Menu Principal",
                        "rows": [
                            {
                                "id": "1",
                                "title": "Fazer uma pergunta",
                                "description": "Envie uma pergunta e eu responderei"
                            },
                            {
                                "id": "2",
                                "title": "Sobre o criador",
                                "description": "Informações sobre o bot"
                            },
                            {
                                "id": "3",
                                "title": "Gerar trabalho",
                                "description": "Gerar um trabalho acadêmico automaticamente"
                            },
                            {
                                "id": "4",
                                "title": "Reclamações",
                                "description": "Envie uma reclamação ou sugestão"
                            },
                            {
                                "id": "5",
                                "title": "Vagas de emprego",
                                "description": "Descubra as vagas mais recentes veja oportunidades de emprego"
                            },
                            {
                                "id": "6",  
                                "title": "Gerar imagens",
                                "description": "Gerar imagens a partir de uma descrição"
                            }
                        ]
                    }
                ]
            }
        }
    }

    response = requests.post(url, headers=headers, json=data)
    logger.info(f"Envio de mensagem de menu: {response.status_code}, {response.text}")
    return response.json()



def send_cta_message(phone_number, button_texts_and_urls):
    phone_number_id = '678628148657980'
    url = f"https://graph.facebook.com/v22.0/{phone_number_id}/messages"
    headers = {
        "Authorization": f"Bearer {WHATSAPP_ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }

    for button_text, button_url in button_texts_and_urls:
        data = {
            "messaging_product": "whatsapp",
            "recipient_type": "individual",
            "to": phone_number,
            "type": "interactive",
            "interactive": {
                "type": "cta_url",
                "header": {
                    "type": "text",
                    "text": "🔍 Vagas de Emprego"
                },
                "body": {
                    "text": "Explore as melhores oportunidades de emprego"
                },
                "footer": {
                    "text": "Nexusbot"
                },
                "action": {
                    "name": "cta_url",
                    "parameters": {
                        "display_text": button_text,
                        "url": button_url
                    }
                }
            }
        }

        response = requests.post(url, headers=headers, json=data)
        logger.info(f"Envio de mensagem CTA: {response.status_code}, {response.text}")

    return response.json()


def send_sticker_message(phone_number, media_id):
    phone_number_id = '678628148657980'
    url = f"https://graph.facebook.com/v22.0/{phone_number_id}/messages"
    headers = {
        "Authorization": f"Bearer {WHATSAPP_ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }

    data = {
        "messaging_product": "whatsapp",
        "recipient_type": "individual",
        "to": phone_number,
        "type": "sticker",
        "sticker": {
            "id": media_id  # ID do sticker que foi carregado previamente
        }
    }

    response = requests.post(url, headers=headers, json=data)
    logger.info(f"Envio de sticker: {response.status_code}, {response.text}")
    return response.json()


# Variável para armazenar o estado da geração do trabalho de cada usuário
user_work_in_progress = {}

def processar_gerar_trabalho(phone_number, tema=None):
    if tema:
        if user_work_in_progress.get(phone_number):
            # Se o trabalho está em progresso, não faz nada
            logger.info(f"O trabalho acadêmico já está em progresso para o número: {phone_number}")
            return

        try:
            # Marca que o trabalho está em progresso
            user_work_in_progress[phone_number] = True
            send_whatsapp_message(phone_number, "🎉 Fantástico! Estamos preparando o seu trabalho. *Por favor, aguarde um momento*  🚀📚")

            # Mensagens de atualização
            mensagens_atualizacao = [
                "⏳ *Estamos trabalhando no seu pedido!* 📚\n\nPor favor, tenha paciência. Seu trabalho está quase pronto!",
                "🔄 *Estamos finalizando os detalhes do seu trabalho!* 💻\n\nAgradecemos pela paciência, está quase pronto!",
                "🕰️ *Estamos quase lá!* 🌟\n\nSó mais um pouco e seu trabalho estará pronto para você!",
                "🚀 *Preparando os últimos ajustes!* 📄\n\nSeu trabalho está a caminho, obrigado por aguardar!",
            ]

            # Envia a mensagem de atualização enquanto o trabalho está sendo gerado
            for i in range(len(mensagens_atualizacao)):  # Envia mensagens diferentes a cada 20 segundos
                time.sleep(20)
                enviar_mensagem_atualizacao(phone_number, i % len(mensagens_atualizacao))

            conteudo = gerar_trabalho_ai21(tema)
            file_path, file_name = criar_trabalho_academico(tema, conteudo)

            if file_path and file_name:
                if phone_number not in file_sent:
                    response = enviar_arquivo_whatsapp(phone_number, file_path, file_name)

                    if response and response.get('status') == 'success':
                        file_sent[phone_number] = file_name
                        # Remover a linha abaixo para evitar o envio do menu duas vezes
                        # send_menu(phone_number)
                    else:
                        logger.error(f"Resposta da API ao enviar o arquivo: {response}")
                else:
                    send_whatsapp_message(phone_number, "O trabalho acadêmico já foi enviado.")

                # Remove o trabalho do estado em progresso
                user_work_in_progress.pop(phone_number, None)
                user_state.pop(phone_number, None)
            else:
                send_whatsapp_message(phone_number, "Erro ao criar o trabalho acadêmico. Tente novamente mais tarde.")
                # Remove o trabalho do estado em progresso
                user_work_in_progress.pop(phone_number, None)
        except Exception as e:
            logger.error(f"Erro ao processar o trabalho: {e}")
            send_whatsapp_message(phone_number, f"Erro ao processar o trabalho: {str(e)}")
            # Remove o trabalho do estado em progresso
            user_work_in_progress.pop(phone_number, None)
            user_state.pop(phone_number, None)


# Processa a escolha do usuário no menu
def process_user_option(phone_number, list_id):
    if list_id == '1':
        send_whatsapp_message(phone_number, "Ótimo! Qual é a sua pergunta?")
        user_state[phone_number] = 'aguardando_pergunta'
    elif list_id == '2':
        send_creator_info(phone_number)
    elif list_id == '3':
        send_whatsapp_message(phone_number, "✨ *Que incrível*! Me diga, sobre qual tema você gostaria que eu escrevesse o seu trabalho?🎓💡")
        user_state[phone_number] = 'aguardando_tema'
    elif list_id == '4':
        send_whatsapp_message(phone_number, "✍🏼 Por favor, descreva sua reclamação ou sugestão e ela será enviada para o administrador.")
        user_state[phone_number] = 'aguardando_reclamacao'
    elif list_id == '5':
        button_texts_and_urls = [
            ("Vagas em Nacionais", "https://www.sovagas.co.mz/"),
            ("Vagas Internacionais", "https://www.linkedin.com/jobs/")
        ]
        send_cta_message(phone_number, button_texts_and_urls)
        # Enviar o sticker após o CTA
        sticker_media_id = '1041485840622792'  # Substitua pelo seu media_id válido
        send_sticker_message(phone_number, sticker_media_id)
    elif list_id == '6':  # Nova opção para gerar imagem com DeepAI
        send_whatsapp_message(phone_number, "Envie a descrição da imagem que deseja gerar:")
        user_state[phone_number] = 'aguardando_descricao_imagem'
    else:
        send_whatsapp_message(phone_number, "Desculpe, não entendi sua escolha. *Por favor, selecione uma opção válida do menu*.")






# Rota para validar o webhook (GET)
@app.route('/webhook', methods=['GET'])
def webhook_verification():
    verify_token = WHATSAPP_ACCESS_TOKEN
    mode = request.args.get('hub.mode')
    token = request.args.get('hub.verify_token')
    challenge = request.args.get('hub.challenge')

    if mode == 'subscribe' and token == verify_token:
        return challenge, 200
    else:
        return 'Erro na verificação do webhook', 403        



# Rota para receber mensagens do WhatsApp (POST)
@app.route('/webhook', methods=['POST'])
def webhook():
    data = request.get_json()
    logger.info(f"Recebendo webhook: {data}")

    if 'entry' in data:
        for entry in data['entry']:
            if 'changes' in entry:
                for change in entry['changes']:
                    if 'value' in change and 'messages' in change['value']:
                        for message in change['value']['messages']:
                            phone_number = message['from']
                            phone_number_id = change['value']['metadata']['phone_number_id']
                            token = WHATSAPP_ACCESS_TOKEN

                            if 'text' in message and 'body' in message['text']:
                                user_message = message['text']['body'].lower()
                                
                                # Verificar se a mensagem é uma saudação usando regex
                                if saudacoes_regex.match(user_message):
                                    send_menu(phone_number)
                                    return jsonify({"status": "success"}), 200

                                elif user_message in ["thanks", "obrigado", "obrigado pela ajuda"]:
                                    send_whatsapp_message(phone_number, "*Conversa encerrada. Até mais!*")
                                    user_state.pop(phone_number, None)
                                    return jsonify({"status": "success"}), 200

                                elif user_state.get(phone_number) == 'aguardando_pergunta':
                                    resposta = perguntar_ai21(user_message)
                                    send_whatsapp_message(phone_number, resposta)
                                    return jsonify({"status": "success"}), 200

                                elif user_state.get(phone_number) == 'aguardando_tema':
                                    processar_gerar_trabalho(phone_number, user_message)
                                    return jsonify({"status": "success"}), 200

                                elif user_state.get(phone_number) == 'aguardando_reclamacao':
                                    admin_phone_number = "+258833390642"
                                    logger.info(f"Enviando reclamação de {phone_number} para o administrador {admin_phone_number}")

                                    response_admin = send_whatsapp_message(admin_phone_number, f"⚠️!! Nova reclamação de {phone_number}:\n\n {user_message}")
                                    response_user = send_whatsapp_message(phone_number, "Sua reclamação foi enviada ao administrador. Obrigado pelo seu feedback!")

                                    logger.info(f"Resposta ao envio da reclamação para o administrador: {response_admin}")
                                    logger.info(f"Resposta ao envio de confirmação ao usuário: {response_user}")
                                    return jsonify({"status": "success"}), 200

                                if user_message in ["imagem", "gerar imagem", "imagem deepai"]:
                                    user_state[phone_number] = 'aguardando_descricao_imagem'
                                    send_whatsapp_message(phone_number, "Envie a descrição da imagem que deseja gerar:")
                                    return jsonify({"status": "success"}), 200

                                elif user_state.get(phone_number) == 'aguardando_descricao_imagem':
                                    handle_image_generation_and_sending(phone_number, user_message, phone_number_id, token)
                                    user_state.pop(phone_number, None)
                                    return jsonify({"status": "success"}), 200

                                # Verifica se há imagem sendo enviada antes de mandar a mensagem padrão
                                if phone_number not in user_state:
                                    send_whatsapp_message(phone_number, "Desculpe, não entendi sua mensagem. *Por favor, escolha uma opção do menu.*")
                                    send_menu(phone_number)
                                    return jsonify({"status": "success"}), 200

                            elif 'interactive' in message and 'list_reply' in message['interactive']:
                                list_id = message['interactive']['list_reply']['id']
                                process_user_option(phone_number, list_id)
                                return jsonify({"status": "success"}), 200  # Retornar após processar a opção da lista

    return jsonify({"status": "success"}), 200  # Resposta padrão se nenhuma condição for atendida




if __name__ == '__main__':
    app.run(port=5000)



